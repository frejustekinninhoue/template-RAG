from __future__ import annotations

import asyncio
import io
import re
from pathlib import Path

import httpx
from bs4 import BeautifulSoup
from docx import Document
from google.oauth2 import service_account
from googleapiclient.discovery import build
from openpyxl import load_workbook


def _normalize_space(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _chunk_text(text: str, chunk_size: int = 2200, overlap: int = 250) -> list[str]:
    cleaned = _normalize_space(text)
    if not cleaned:
        return []

    chunks: list[str] = []
    start = 0
    while start < len(cleaned):
        end = min(start + chunk_size, len(cleaned))
        chunks.append(cleaned[start:end])
        if end == len(cleaned):
            break
        start = max(0, end - overlap)
    return chunks


def build_tree_from_text(title: str, text: str, source_type: str) -> list[dict]:
    chunks = _chunk_text(text)
    if not chunks:
        return []

    tree: list[dict] = []
    for idx, chunk in enumerate(chunks, start=1):
        summary = chunk[:300] + ("..." if len(chunk) > 300 else "")
        tree.append(
            {
                "node_id": f"{source_type}-{idx}",
                "title": f"{title} - part {idx}",
                "summary": summary,
                "text": chunk,
            }
        )
    return tree


def parse_docx_text(path: str) -> str:
    doc = Document(path)

    lines: list[str] = []
    for p in doc.paragraphs:
        value = p.text.strip()
        if value:
            lines.append(value)

    for table in doc.tables:
        for row in table.rows:
            cells = [_normalize_space(c.text) for c in row.cells]
            row_text = " | ".join([c for c in cells if c])
            if row_text:
                lines.append(row_text)

    return "\n".join(lines)


def parse_xlsx_text(path: str) -> str:
    wb = load_workbook(path, data_only=True)
    blocks: list[str] = []

    for sheet in wb.worksheets:
        blocks.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = [str(v).strip() for v in row if v is not None and str(v).strip()]
            if values:
                blocks.append(" | ".join(values))

    return "\n".join(blocks)


async def fetch_url_text(url: str, timeout: float = 25.0) -> str:
    async with httpx.AsyncClient(follow_redirects=True, timeout=timeout) as client:
        resp = await client.get(url)
        resp.raise_for_status()

    soup = BeautifulSoup(resp.text, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    text = soup.get_text(separator="\n")
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return "\n".join(lines)


def _google_credentials(service_account_file: str, scopes: list[str]):
    return service_account.Credentials.from_service_account_file(
        service_account_file,
        scopes=scopes,
    )


def fetch_google_doc_text(doc_id: str, service_account_file: str) -> str:
    creds = _google_credentials(
        service_account_file,
        ["https://www.googleapis.com/auth/drive.readonly"],
    )
    drive = build("drive", "v3", credentials=creds, cache_discovery=False)

    request = drive.files().export_media(fileId=doc_id, mimeType="text/plain")
    content = request.execute()
    if isinstance(content, bytes):
        return content.decode("utf-8", errors="ignore")
    return str(content)


def fetch_google_sheet_text(sheet_id: str, service_account_file: str) -> str:
    creds = _google_credentials(
        service_account_file,
        ["https://www.googleapis.com/auth/spreadsheets.readonly"],
    )
    sheets = build("sheets", "v4", credentials=creds, cache_discovery=False)

    meta = sheets.spreadsheets().get(spreadsheetId=sheet_id).execute()
    sheet_names = [s["properties"]["title"] for s in meta.get("sheets", [])]

    lines: list[str] = []
    for name in sheet_names:
        lines.append(f"# Sheet: {name}")
        result = sheets.spreadsheets().values().get(
            spreadsheetId=sheet_id,
            range=f"'{name}'",
        ).execute()

        for row in result.get("values", []):
            values = [str(v).strip() for v in row if str(v).strip()]
            if values:
                lines.append(" | ".join(values))

    return "\n".join(lines)


async def text_from_google_doc(doc_id: str, service_account_file: str) -> str:
    return await asyncio.to_thread(fetch_google_doc_text, doc_id, service_account_file)


async def text_from_google_sheet(sheet_id: str, service_account_file: str) -> str:
    return await asyncio.to_thread(fetch_google_sheet_text, sheet_id, service_account_file)
